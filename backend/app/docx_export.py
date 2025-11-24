# backend/app/docx_export.py
from io import BytesIO
from docx import Document
from .models import Project

def build_docx(project: Project) -> BytesIO:
    doc = Document()
    doc.add_heading(project.title, level=0)
    doc.add_paragraph(f"Topic: {project.topic}")

    for section in sorted(project.sections, key=lambda s: s.order):
        doc.add_heading(section.title, level=1)
        doc.add_paragraph(section.content or "")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
